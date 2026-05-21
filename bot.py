import os
import copy
import logging
import threading
import base64
import json
from datetime import datetime
from http.server import HTTPServer, BaseHTTPRequestHandler
from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup
from telegram.ext import (
    Application, CommandHandler, MessageHandler,
    CallbackQueryHandler, ContextTypes, filters, ConversationHandler
)
from openpyxl import load_workbook
from docx import Document
import httpx

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

BOT_TOKEN = os.environ.get("BOT_TOKEN", "")
SEP_CHAT_ID = int(os.environ.get("SEP_CHAT_ID", "0"))
NHOM_VAT_TU_ID = int(os.environ.get("NHOM_VAT_TU_ID", "0"))
ANTHROPIC_API_KEY = os.environ.get("ANTHROPIC_API_KEY", "")

# ====== HEALTH SERVER ======
class HealthHandler(BaseHTTPRequestHandler):
    def do_GET(self):
        self.send_response(200)
        self.end_headers()
        self.wfile.write(b'Bot dang chay!')
    def log_message(self, format, *args):
        pass

def run_health_server():
    port = int(os.environ.get("PORT", 8080))
    server = HTTPServer(("0.0.0.0", port), HealthHandler)
    server.serve_forever()

# ====== TRẠNG THÁI ======
(VT_BO_PHAN, VT_MA, VT_TEN, VT_DVT, VT_NHU_CAU,
 VT_TON_KHO, VT_SO_LUONG, VT_XUAT_XU, VT_MUC_DICH,
 VT_NGAY_GIAO, VT_GHI_CHU, VT_THEM, VT_XAC_NHAN) = range(13)
(DX_BO_PHAN, DX_NOI_DUNG, DX_XAC_NHAN) = range(13, 16)
(TT_CHUC_DANH, TT_BO_PHAN, TT_NOI_DUNG, TT_SO_HD, TT_NGAY_HD,
 TT_SO_TIEN, TT_HINH_THUC, TT_SO_TK, TT_NGAN_HANG,
 TT_CHU_TK, TT_XAC_NHAN) = range(16, 27)
(EXCEL_BO_PHAN, EXCEL_XAC_NHAN) = range(27, 29)
(HINH_BO_PHAN, HINH_XAC_NHAN) = range(29, 31)

store = {}

# ====== MENU CHÍNH ======
async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    keyboard = [
        [InlineKeyboardButton("📦 Yêu cầu vật tư (nhắn tin)", callback_data="menu_vattu")],
        [InlineKeyboardButton("📊 Yêu cầu vật tư (upload Excel)", callback_data="menu_excel")],
        [InlineKeyboardButton("📸 Yêu cầu vật tư (chụp hình)", callback_data="menu_hinh")],
        [InlineKeyboardButton("📝 Phiếu đề xuất", callback_data="menu_deuxuat")],
        [InlineKeyboardButton("💰 Đề nghị thanh toán", callback_data="menu_thanhtoan")],
    ]
    await update.message.reply_text(
        "👋 Xin chào! Tôi là *Bot Công ty CORNER*\n\nBạn muốn làm gì?",
        reply_markup=InlineKeyboardMarkup(keyboard),
        parse_mode="Markdown"
    )

# ====== LUỒNG UPLOAD EXCEL ======
async def excel_start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user = update.effective_user
    uid = user.id
    store[f"excel_{uid}"] = {
        "id": f"excel_{uid}_{int(datetime.now().timestamp())}",
        "nguoi_de_nghi": f"{user.first_name} {user.last_name or ''}".strip(),
        "bo_phan": "", "vat_tu": []
    }
    msg = (
        "📊 *UPLOAD FILE EXCEL*\n\n"
        "Bạn thuộc Phòng/Bộ phận nào?\n"
        "_(Sau đó gửi file Excel có danh sách vật tư)_"
    )
    if update.callback_query:
        await update.callback_query.answer()
        await update.callback_query.message.reply_text(msg, parse_mode="Markdown")
    else:
        await update.message.reply_text(msg, parse_mode="Markdown")
    return EXCEL_BO_PHAN

async def excel_bo_phan(update, context):
    uid = update.effective_user.id
    store[f"excel_{uid}"]["bo_phan"] = update.message.text
    await update.message.reply_text(
        "✅ Bây giờ hãy gửi file Excel lên!\n\n"
        "📋 *Yêu cầu file Excel:*\n"
        "Cột A: Mã vật tư\n"
        "Cột B: Tên vật tư\n"
        "Cột C: Đơn vị tính\n"
        "Cột D: Số lượng cần mua\n"
        "Cột E: Mục đích sử dụng\n"
        "_(Dòng 1 là tiêu đề, dữ liệu từ dòng 2)_",
        parse_mode="Markdown"
    )
    return EXCEL_XAC_NHAN

async def excel_nhan_file(update, context):
    uid = update.effective_user.id
    if f"excel_{uid}" not in store:
        await update.message.reply_text("⚠️ Vui lòng gõ /start và chọn lại!")
        return ConversationHandler.END

    await update.message.reply_text("⏳ Đang đọc file Excel...")
    try:
        file = await update.message.document.get_file()
        ten_file_tai = f"upload_{uid}.xlsx"
        await file.download_to_drive(ten_file_tai)

        wb = load_workbook(ten_file_tai)
        ws = wb.active
        vat_tu_list = []
        for row in ws.iter_rows(min_row=2, values_only=True):
            if not row[1]:
                continue
            vat_tu_list.append({
                "ma_vat_tu": str(row[0] or ""),
                "ten_vat_tu": str(row[1] or ""),
                "dvt": str(row[2] or ""),
                "so_luong_mua": str(row[3] or ""),
                "muc_dich": str(row[4] or ""),
                "nhu_cau": str(row[3] or ""),
                "ton_kho": "0",
                "xuat_xu": "",
                "ngay_giao": "",
                "ghi_chu": "",
            })

        if not vat_tu_list:
            await update.message.reply_text("❌ File Excel trống hoặc sai định dạng!")
            return EXCEL_XAC_NHAN

        store[f"excel_{uid}"]["vat_tu"] = vat_tu_list
        d = store[f"excel_{uid}"]

        lines = [f"📋 *ĐÃ ĐỌC ĐƯỢC {len(vat_tu_list)} VẬT TƯ:*\n🏢 {d['bo_phan']}\n"]
        for i, vt in enumerate(vat_tu_list[:10], 1):
            lines.append(f"{i}. {vt['ten_vat_tu']} — {vt['so_luong_mua']} {vt['dvt']}")
        if len(vat_tu_list) > 10:
            lines.append(f"... và {len(vat_tu_list)-10} vật tư khác")

        keyboard = [[
            InlineKeyboardButton("✅ Xác nhận gửi sếp", callback_data=f"excel_ok_{uid}"),
            InlineKeyboardButton("❌ Hủy", callback_data=f"excel_huy_{uid}"),
        ]]
        await update.message.reply_text(
            "\n".join(lines),
            reply_markup=InlineKeyboardMarkup(keyboard),
            parse_mode="Markdown"
        )
        return EXCEL_XAC_NHAN

    except Exception as e:
        logger.error(e)
        await update.message.reply_text("❌ Lỗi đọc file. Hãy kiểm tra định dạng và thử lại!")
        return EXCEL_XAC_NHAN

async def excel_xacnhan_callback(update, context):
    query = update.callback_query
    await query.answer()
    uid = query.from_user.id
    if query.data.startswith("excel_huy"):
        store.pop(f"excel_{uid}", None)
        await query.message.reply_text("❌ Đã hủy.")
        return ConversationHandler.END
    d = store[f"excel_{uid}"]
    await query.message.reply_text("⏳ Đang tạo phiếu và gửi sếp...")
    try:
        ten_file = tao_phieu_vattu(d)
        tom_tat = f"📦 *YÊU CẦU VẬT TƯ (Excel)*\n👤 {d['nguoi_de_nghi']} | 🏢 {d['bo_phan']}\n"
        for i, vt in enumerate(d["vat_tu"][:5], 1):
            tom_tat += f"\n{i}. {vt['ten_vat_tu']} — {vt['so_luong_mua']} {vt['dvt']}"
        if len(d["vat_tu"]) > 5:
            tom_tat += f"\n... và {len(d['vat_tu'])-5} vật tư khác"
        keyboard = [[
            InlineKeyboardButton("✅ DUYỆT", callback_data=f"duyet_vt_{d['id']}"),
            InlineKeyboardButton("❌ TỪ CHỐI", callback_data=f"tuchoi_vt_{d['id']}"),
        ]]
        with open(ten_file, "rb") as f:
            await context.bot.send_document(
                chat_id=SEP_CHAT_ID, document=f, filename=ten_file,
                caption=tom_tat + "\n\n👆 Vui lòng duyệt hoặc từ chối.",
                reply_markup=InlineKeyboardMarkup(keyboard), parse_mode="Markdown"
            )
        store[f"don_vt_{d['id']}"] = {"don": d, "uid": uid, "file": ten_file}
        await query.message.reply_text("✅ Đã gửi sếp duyệt!")
    except Exception as e:
        logger.error(e)
        await query.message.reply_text("❌ Có lỗi xảy ra.")
    store.pop(f"excel_{uid}", None)
    return ConversationHandler.END

# ====== LUỒNG CHỤP HÌNH ======
async def hinh_start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user = update.effective_user
    uid = user.id
    store[f"hinh_{uid}"] = {
        "id": f"hinh_{uid}_{int(datetime.now().timestamp())}",
        "nguoi_de_nghi": f"{user.first_name} {user.last_name or ''}".strip(),
        "bo_phan": "", "vat_tu": []
    }
    msg = "📸 *CHỤP HÌNH PHIẾU VẬT TƯ*\n\nBạn thuộc Phòng/Bộ phận nào?"
    if update.callback_query:
        await update.callback_query.answer()
        await update.callback_query.message.reply_text(msg, parse_mode="Markdown")
    else:
        await update.message.reply_text(msg, parse_mode="Markdown")
    return HINH_BO_PHAN

async def hinh_bo_phan(update, context):
    uid = update.effective_user.id
    store[f"hinh_{uid}"]["bo_phan"] = update.message.text
    await update.message.reply_text(
        "✅ Bây giờ hãy chụp hình phiếu vật tư và gửi lên!\n\n"
        "📸 *Lưu ý:*\n"
        "• Chụp rõ nét, đủ ánh sáng\n"
        "• Đảm bảo chữ đọc được\n"
        "• Có thể chụp nhiều trang",
        parse_mode="Markdown"
    )
    return HINH_XAC_NHAN

async def hinh_nhan_anh(update, context):
    uid = update.effective_user.id
    if f"hinh_{uid}" not in store:
        await update.message.reply_text("⚠️ Vui lòng gõ /start và chọn lại!")
        return ConversationHandler.END

    await update.message.reply_text("⏳ Đang dùng AI đọc hình ảnh...")
    try:
        photo = update.message.photo[-1]
        file = await photo.get_file()
        ten_file_anh = f"anh_{uid}.jpg"
        await file.download_to_drive(ten_file_anh)

        with open(ten_file_anh, "rb") as f:
            anh_base64 = base64.b64encode(f.read()).decode()

        headers = {
            "x-api-key": ANTHROPIC_API_KEY,
            "anthropic-version": "2023-06-01",
            "content-type": "application/json"
        }
        payload = {
            "model": "claude-opus-4-20250514",
            "max_tokens": 2000,
            "messages": [{
                "role": "user",
                "content": [
                    {
                        "type": "image",
                        "source": {
                            "type": "base64",
                            "media_type": "image/jpeg",
                            "data": anh_base64
                        }
                    },
                    {
                        "type": "text",
                        "text": """Đọc hình ảnh phiếu yêu cầu vật tư này và trả về JSON với format sau (không có markdown):
{
  "vat_tu": [
    {
      "ma_vat_tu": "",
      "ten_vat_tu": "",
      "dvt": "",
      "so_luong_mua": "",
      "muc_dich": "",
      "nhu_cau": "",
      "ton_kho": "0",
      "xuat_xu": "",
      "ngay_giao": "",
      "ghi_chu": ""
    }
  ]
}
Chỉ trả về JSON, không có text khác."""
                    }
                ]
            }]
        }

        async with httpx.AsyncClient(timeout=30) as client:
            resp = await client.post(
                "https://api.anthropic.com/v1/messages",
                headers=headers,
                json=payload
            )
            result = resp.json()

        text = result["content"][0]["text"].strip()
        data = json.loads(text)
        vat_tu_list = data.get("vat_tu", [])

        if not vat_tu_list:
            await update.message.reply_text("❌ AI không đọc được vật tư. Hãy chụp lại rõ hơn!")
            return HINH_XAC_NHAN

        store[f"hinh_{uid}"]["vat_tu"] = vat_tu_list
        d = store[f"hinh_{uid}"]

        lines = [f"📋 *AI ĐỌC ĐƯỢC {len(vat_tu_list)} VẬT TƯ:*\n🏢 {d['bo_phan']}\n"]
        for i, vt in enumerate(vat_tu_list, 1):
            lines.append(f"{i}. {vt['ten_vat_tu']} — {vt['so_luong_mua']} {vt['dvt']}")

        keyboard = [[
            InlineKeyboardButton("✅ Xác nhận gửi sếp", callback_data=f"hinh_ok_{uid}"),
            InlineKeyboardButton("❌ Hủy", callback_data=f"hinh_huy_{uid}"),
        ]]
        await update.message.reply_text(
            "\n".join(lines),
            reply_markup=InlineKeyboardMarkup(keyboard),
            parse_mode="Markdown"
        )
        return HINH_XAC_NHAN

    except Exception as e:
        logger.error(e)
        await update.message.reply_text("❌ Lỗi đọc hình. Thử lại hoặc dùng cách upload Excel!")
        return HINH_XAC_NHAN

async def hinh_xacnhan_callback(update, context):
    query = update.callback_query
    await query.answer()
    uid = query.from_user.id
    if query.data.startswith("hinh_huy"):
        store.pop(f"hinh_{uid}", None)
        await query.message.reply_text("❌ Đã hủy.")
        return ConversationHandler.END
    d = store[f"hinh_{uid}"]
    await query.message.reply_text("⏳ Đang tạo phiếu và gửi sếp...")
    try:
        ten_file = tao_phieu_vattu(d)
        tom_tat = f"📦 *YÊU CẦU VẬT TƯ (Chụp hình)*\n👤 {d['nguoi_de_nghi']} | 🏢 {d['bo_phan']}\n"
        for i, vt in enumerate(d["vat_tu"][:5], 1):
            tom_tat += f"\n{i}. {vt['ten_vat_tu']} — {vt['so_luong_mua']} {vt['dvt']}"
        if len(d["vat_tu"]) > 5:
            tom_tat += f"\n... và {len(d['vat_tu'])-5} vật tư khác"
        keyboard = [[
            InlineKeyboardButton("✅ DUYỆT", callback_data=f"duyet_vt_{d['id']}"),
            InlineKeyboardButton("❌ TỪ CHỐI", callback_data=f"tuchoi_vt_{d['id']}"),
        ]]
        with open(ten_file, "rb") as f:
            await context.bot.send_document(
                chat_id=SEP_CHAT_ID, document=f, filename=ten_file,
                caption=tom_tat + "\n\n👆 Vui lòng duyệt hoặc từ chối.",
                reply_markup=InlineKeyboardMarkup(keyboard), parse_mode="Markdown"
            )
        store[f"don_vt_{d['id']}"] = {"don": d, "uid": uid, "file": ten_file}
        await query.message.reply_text("✅ Đã gửi sếp duyệt!")
    except Exception as e:
        logger.error(e)
        await query.message.reply_text("❌ Có lỗi xảy ra.")
    store.pop(f"hinh_{uid}", None)
    return ConversationHandler.END

# ====== LUỒNG VẬT TƯ NHẮN TIN ======
async def vattu_start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user = update.effective_user
    uid = user.id
    store[uid] = {
        "id": f"{uid}_{int(datetime.now().timestamp())}",
        "nguoi_de_nghi": f"{user.first_name} {user.last_name or ''}".strip(),
        "bo_phan": "", "vat_tu": [], "vat_tu_hien_tai": {}
    }
    msg = "📦 *YÊU CẦU VẬT TƯ*\n\nBạn thuộc Phòng/Bộ phận nào?"
    if update.callback_query:
        await update.callback_query.answer()
        await update.callback_query.message.reply_text(msg, parse_mode="Markdown")
    else:
        await update.message.reply_text(msg, parse_mode="Markdown")
    return VT_BO_PHAN

async def vt_bo_phan(update, context):
    uid = update.effective_user.id
    store[uid]["bo_phan"] = update.message.text
    await update.message.reply_text("Mã vật tư (gõ `-` nếu không có):")
    return VT_MA

async def vt_ma(update, context):
    uid = update.effective_user.id
    v = update.message.text.strip()
    store[uid]["vat_tu_hien_tai"]["ma_vat_tu"] = "" if v == "-" else v
    await update.message.reply_text("Tên vật tư:")
    return VT_TEN

async def vt_ten(update, context):
    uid = update.effective_user.id
    store[uid]["vat_tu_hien_tai"]["ten_vat_tu"] = update.message.text
    await update.message.reply_text("Đơn vị tính (cái, kg, m...):")
    return VT_DVT

async def vt_dvt(update, context):
    uid = update.effective_user.id
    store[uid]["vat_tu_hien_tai"]["dvt"] = update.message.text
    await update.message.reply_text("Số lượng cần mua:")
    return VT_SO_LUONG

async def vt_so_luong(update, context):
    uid = update.effective_user.id
    store[uid]["vat_tu_hien_tai"]["so_luong_mua"] = update.message.text
    store[uid]["vat_tu_hien_tai"]["nhu_cau"] = update.message.text
    store[uid]["vat_tu_hien_tai"]["ton_kho"] = "0"
    store[uid]["vat_tu_hien_tai"]["xuat_xu"] = ""
    store[uid]["vat_tu_hien_tai"]["ngay_giao"] = ""
    store[uid]["vat_tu_hien_tai"]["ghi_chu"] = ""
    await update.message.reply_text("Mục đích sử dụng:")
    return VT_MUC_DICH

async def vt_muc_dich(update, context):
    uid = update.effective_user.id
    store[uid]["vat_tu_hien_tai"]["muc_dich"] = update.message.text
    store[uid]["vat_tu"].append(copy.deepcopy(store[uid]["vat_tu_hien_tai"]))
    store[uid]["vat_tu_hien_tai"] = {}
    so = len(store[uid]["vat_tu"])
    keyboard = [
        [InlineKeyboardButton("➕ Thêm vật tư nữa", callback_data="vt_them")],
        [InlineKeyboardButton("✅ Xong, gửi duyệt", callback_data="vt_gui")],
    ]
    await update.message.reply_text(
        f"✅ Đã thêm vật tư thứ *{so}*! Tiếp theo?",
        reply_markup=InlineKeyboardMarkup(keyboard),
        parse_mode="Markdown"
    )
    return VT_THEM

async def vt_them_callback(update, context):
    query = update.callback_query
    await query.answer()
    uid = query.from_user.id
    if query.data == "vt_them":
        await query.message.reply_text("Mã vật tư tiếp theo (gõ `-` nếu không có):")
        return VT_MA
    elif query.data == "vt_gui":
        d = store[uid]
        lines = [f"📋 *YÊU CẦU VẬT TƯ*\n👤 {d['nguoi_de_nghi']} | 🏢 {d['bo_phan']}\n"]
        for i, vt in enumerate(d["vat_tu"], 1):
            lines.append(f"*{i}. {vt['ten_vat_tu']}* — {vt['so_luong_mua']} {vt['dvt']}")
        keyboard = [
            [InlineKeyboardButton("✅ Xác nhận gửi", callback_data="vt_xacnhan")],
            [InlineKeyboardButton("❌ Hủy", callback_data="vt_huy")],
        ]
        await query.message.reply_text(
            "\n".join(lines),
            reply_markup=InlineKeyboardMarkup(keyboard),
            parse_mode="Markdown"
        )
        return VT_XAC_NHAN

async def vt_xacnhan_callback(update, context):
    query = update.callback_query
    await query.answer()
    uid = query.from_user.id
    if query.data == "vt_huy":
        store.pop(uid, None)
        await query.message.reply_text("❌ Đã hủy.")
        return ConversationHandler.END
    d = store[uid]
    await query.message.reply_text("⏳ Đang tạo phiếu và gửi sếp...")
    try:
        ten_file = tao_phieu_vattu(d)
        tom_tat = f"📦 *YÊU CẦU VẬT TƯ*\n👤 {d['nguoi_de_nghi']} | 🏢 {d['bo_phan']}\n"
        for i, vt in enumerate(d["vat_tu"], 1):
            tom_tat += f"\n{i}. {vt['ten_vat_tu']} — {vt['so_luong_mua']} {vt['dvt']}"
        keyboard = [[
            InlineKeyboardButton("✅ DUYỆT", callback_data=f"duyet_vt_{d['id']}"),
            InlineKeyboardButton("❌ TỪ CHỐI", callback_data=f"tuchoi_vt_{d['id']}"),
        ]]
        with open(ten_file, "rb") as f:
            await context.bot.send_document(
                chat_id=SEP_CHAT_ID, document=f, filename=ten_file,
                caption=tom_tat + "\n\n👆 Vui lòng duyệt hoặc từ chối.",
                reply_markup=InlineKeyboardMarkup(keyboard), parse_mode="Markdown"
            )
        store[f"don_vt_{d['id']}"] = {"don": d, "uid": uid, "file": ten_file}
        await query.message.reply_text("✅ Đã gửi sếp duyệt!")
    except Exception as e:
        logger.error(e)
        await query.message.reply_text("❌ Có lỗi xảy ra.")
    store.pop(uid, None)
    return ConversationHandler.END

# ====== LUỒNG ĐỀ XUẤT ======
async def deuxuat_start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user = update.effective_user
    uid = user.id
    store[f"dx_{uid}"] = {
        "id": f"dx_{uid}_{int(datetime.now().timestamp())}",
        "ho_ten": f"{user.first_name} {user.last_name or ''}".strip(),
        "bo_phan": "", "noi_dung": ""
    }
    msg = "📝 *PHIẾU ĐỀ XUẤT*\n\nBạn thuộc Phòng/Bộ phận nào?"
    if update.callback_query:
        await update.callback_query.answer()
        await update.callback_query.message.reply_text(msg, parse_mode="Markdown")
    else:
        await update.message.reply_text(msg, parse_mode="Markdown")
    return DX_BO_PHAN

async def dx_bo_phan(update, context):
    uid = update.effective_user.id
    store[f"dx_{uid}"]["bo_phan"] = update.message.text
    await update.message.reply_text("📋 Nội dung đề xuất:", parse_mode="Markdown")
    return DX_NOI_DUNG

async def dx_noi_dung(update, context):
    uid = update.effective_user.id
    store[f"dx_{uid}"]["noi_dung"] = update.message.text
    d = store[f"dx_{uid}"]
    keyboard = [
        [InlineKeyboardButton("✅ Xác nhận gửi", callback_data=f"dx_xacnhan_{uid}")],
        [InlineKeyboardButton("❌ Hủy", callback_data=f"dx_huy_{uid}")],
    ]
    await update.message.reply_text(
        f"📝 *XEM LẠI:*\n👤 {d['ho_ten']} | 🏢 {d['bo_phan']}\n\n{d['noi_dung']}",
        reply_markup=InlineKeyboardMarkup(keyboard),
        parse_mode="Markdown"
    )
    return DX_XAC_NHAN

async def dx_xacnhan_callback(update, context):
    query = update.callback_query
    await query.answer()
    uid = query.from_user.id
    if query.data.startswith("dx_huy"):
        store.pop(f"dx_{uid}", None)
        await query.message.reply_text("❌ Đã hủy.")
        return ConversationHandler.END
    d = store[f"dx_{uid}"]
    try:
        ten_file = tao_phieu_deuxuat(d)
        tom_tat = f"📝 *PHIẾU ĐỀ XUẤT*\n👤 {d['ho_ten']} | 🏢 {d['bo_phan']}\n\n{d['noi_dung']}"
        keyboard = [[
            InlineKeyboardButton("✅ DUYỆT", callback_data=f"duyet_dx_{d['id']}"),
            InlineKeyboardButton("❌ TỪ CHỐI", callback_data=f"tuchoi_dx_{d['id']}"),
        ]]
        with open(ten_file, "rb") as f:
            await context.bot.send_document(
                chat_id=SEP_CHAT_ID, document=f, filename=ten_file,
                caption=tom_tat + "\n\n👆 Vui lòng duyệt hoặc từ chối.",
                reply_markup=InlineKeyboardMarkup(keyboard), parse_mode="Markdown"
            )
        store[f"don_dx_{d['id']}"] = {"don": d, "uid": uid, "file": ten_file}
        await query.message.reply_text("✅ Đã gửi sếp duyệt!")
    except Exception as e:
        logger.error(e)
        await query.message.reply_text("❌ Có lỗi xảy ra.")
    store.pop(f"dx_{uid}", None)
    return ConversationHandler.END

# ====== LUỒNG THANH TOÁN ======
async def thanhtoan_start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user = update.effective_user
    uid = user.id
    store[f"tt_{uid}"] = {
        "id": f"tt_{uid}_{int(datetime.now().timestamp())}",
        "ho_ten": f"{user.first_name} {user.last_name or ''}".strip(),
        "chuc_danh": "", "bo_phan": "", "noi_dung": "",
        "so_hd": "", "ngay_hd": "", "so_tien": "",
        "hinh_thuc": "", "so_tk": "", "ngan_hang": "", "chu_tk": ""
    }
    msg = "💰 *ĐỀ NGHỊ THANH TOÁN*\n\nChức danh của bạn?"
    if update.callback_query:
        await update.callback_query.answer()
        await update.callback_query.message.reply_text(msg, parse_mode="Markdown")
    else:
        await update.message.reply_text(msg, parse_mode="Markdown")
    return TT_CHUC_DANH

async def tt_chuc_danh(update, context):
    uid = update.effective_user.id
    store[f"tt_{uid}"]["chuc_danh"] = update.message.text
    await update.message.reply_text("Phòng/Bộ phận:")
    return TT_BO_PHAN

async def tt_bo_phan(update, context):
    uid = update.effective_user.id
    store[f"tt_{uid}"]["bo_phan"] = update.message.text
    await update.message.reply_text("Nội dung thanh toán:")
    return TT_NOI_DUNG

async def tt_noi_dung(update, context):
    uid = update.effective_user.id
    store[f"tt_{uid}"]["noi_dung"] = update.message.text
    await update.message.reply_text("Số hóa đơn (gõ `-` nếu không có):")
    return TT_SO_HD

async def tt_so_hd(update, context):
    uid = update.effective_user.id
    v = update.message.text.strip()
    store[f"tt_{uid}"]["so_hd"] = "" if v == "-" else v
    await update.message.reply_text("Ngày hóa đơn (gõ `-` nếu không có):")
    return TT_NGAY_HD

async def tt_ngay_hd(update, context):
    uid = update.effective_user.id
    v = update.message.text.strip()
    store[f"tt_{uid}"]["ngay_hd"] = "" if v == "-" else v
    await update.message.reply_text("Số tiền:")
    return TT_SO_TIEN

async def tt_so_tien(update, context):
    uid = update.effective_user.id
    store[f"tt_{uid}"]["so_tien"] = update.message.text
    keyboard = [
        [InlineKeyboardButton("💵 Tiền mặt", callback_data="tt_ht_tienmat")],
        [InlineKeyboardButton("🏦 Chuyển khoản", callback_data="tt_ht_ck")],
    ]
    await update.message.reply_text("Hình thức thanh toán:", reply_markup=InlineKeyboardMarkup(keyboard))
    return TT_HINH_THUC

async def tt_hinh_thuc_callback(update, context):
    query = update.callback_query
    await query.answer()
    uid = query.from_user.id
    ht = "Tiền mặt" if query.data == "tt_ht_tienmat" else "Chuyển khoản"
    store[f"tt_{uid}"]["hinh_thuc"] = ht
    if ht == "Tiền mặt":
        d = store[f"tt_{uid}"]
        keyboard = [[
            InlineKeyboardButton("✅ Xác nhận gửi", callback_data=f"tt_xacnhan_{uid}"),
            InlineKeyboardButton("❌ Hủy", callback_data=f"tt_huy_{uid}"),
        ]]
        await query.message.reply_text(
            f"💰 *XEM LẠI:*\n👤 {d['ho_ten']} — {d['chuc_danh']} | 🏢 {d['bo_phan']}\n"
            f"📋 {d['noi_dung']}\n💵 {d['so_tien']} VNĐ — Tiền mặt",
            reply_markup=InlineKeyboardMarkup(keyboard), parse_mode="Markdown"
        )
        return TT_XAC_NHAN
    else:
        await query.message.reply_text("Số tài khoản:")
        return TT_SO_TK

async def tt_so_tk(update, context):
    uid = update.effective_user.id
    store[f"tt_{uid}"]["so_tk"] = update.message.text
    await update.message.reply_text("Ngân hàng:")
    return TT_NGAN_HANG

async def tt_ngan_hang(update, context):
    uid = update.effective_user.id
    store[f"tt_{uid}"]["ngan_hang"] = update.message.text
    await update.message.reply_text("Chủ tài khoản:")
    return TT_CHU_TK

async def tt_chu_tk(update, context):
    uid = update.effective_user.id
    store[f"tt_{uid}"]["chu_tk"] = update.message.text
    d = store[f"tt_{uid}"]
    keyboard = [[
        InlineKeyboardButton("✅ Xác nhận gửi", callback_data=f"tt_xacnhan_{uid}"),
        InlineKeyboardButton("❌ Hủy", callback_data=f"tt_huy_{uid}"),
    ]]
    await update.message.reply_text(
        f"💰 *XEM LẠI:*\n👤 {d['ho_ten']} — {d['chuc_danh']} | 🏢 {d['bo_phan']}\n"
        f"📋 {d['noi_dung']}\n💵 {d['so_tien']} VNĐ — {d['hinh_thuc']}\n"
        f"🏦 {d['so_tk']} | {d['ngan_hang']} | {d['chu_tk']}",
        reply_markup=InlineKeyboardMarkup(keyboard), parse_mode="Markdown"
    )
    return TT_XAC_NHAN

async def tt_xacnhan_callback(update, context):
    query = update.callback_query
    await query.answer()
    uid = query.from_user.id
    if query.data.startswith("tt_huy"):
        store.pop(f"tt_{uid}", None)
        await query.message.reply_text("❌ Đã hủy.")
        return ConversationHandler.END
    d = store[f"tt_{uid}"]
    try:
        ten_file = tao_phieu_thanhtoan(d)
        tom_tat = (f"💰 *ĐỀ NGHỊ THANH TOÁN*\n👤 {d['ho_ten']} — {d['chuc_danh']} | 🏢 {d['bo_phan']}\n"
                   f"📋 {d['noi_dung']}\n💵 {d['so_tien']} VNĐ — {d['hinh_thuc']}")
        keyboard = [[
            InlineKeyboardButton("✅ DUYỆT", callback_data=f"duyet_tt_{d['id']}"),
            InlineKeyboardButton("❌ TỪ CHỐI", callback_data=f"tuchoi_tt_{d['id']}"),
        ]]
        with open(ten_file, "rb") as f:
            await context.bot.send_document(
                chat_id=SEP_CHAT_ID, document=f, filename=ten_file,
                caption=tom_tat + "\n\n👆 Vui lòng duyệt hoặc từ chối.",
                reply_markup=InlineKeyboardMarkup(keyboard), parse_mode="Markdown"
            )
        store[f"don_tt_{d['id']}"] = {"don": d, "uid": uid, "file": ten_file}
        await query.message.reply_text("✅ Đã gửi sếp duyệt!")
    except Exception as e:
        logger.error(e)
        await query.message.reply_text("❌ Có lỗi xảy ra.")
    store.pop(f"tt_{uid}", None)
    return ConversationHandler.END

# ====== XỬ LÝ DUYỆT ======
async def xu_ly_duyet(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()
    data = query.data
    trang_thai = "duyet" if data.startswith("duyet_") else "tuchoi"
    loai_id = data.replace("duyet_", "").replace("tuchoi_", "")
    key = f"don_{loai_id}"
    if key not in store:
        await query.message.reply_text("⚠️ Không tìm thấy phiếu này.")
        return
    info = store[key]
    d = info["don"]
    uid = info["uid"]
    if trang_thai == "duyet":
        if loai_id.startswith("vt_"):
            tom_tat = f"📦 *YÊU CẦU VẬT TƯ ĐÃ DUYỆT*\n👤 {d.get('nguoi_de_nghi', d.get('ho_ten',''))} | 🏢 {d['bo_phan']}\n"
            for i, vt in enumerate(d["vat_tu"], 1):
                tom_tat += f"\n{i}. {vt['ten_vat_tu']} — {vt['so_luong_mua']} {vt['dvt']}"
            with open(info["file"], "rb") as f:
                await context.bot.send_document(
                    chat_id=NHOM_VAT_TU_ID, document=f, filename=info["file"],
                    caption=tom_tat + "\n\n👉 Tiến hành mua vật tư.",
                    parse_mode="Markdown"
                )
        await context.bot.send_message(chat_id=uid, text="✅ *Phiếu của bạn đã được DUYỆT!*", parse_mode="Markdown")
        try:
            await query.message.edit_caption(caption=query.message.caption + "\n\n✅ *ĐÃ DUYỆT*", parse_mode="Markdown")
        except: pass
    else:
        await context.bot.send_message(chat_id=uid, text="❌ *Phiếu của bạn đã bị TỪ CHỐI.*", parse_mode="Markdown")
        try:
            await query.message.edit_caption(caption=query.message.caption + "\n\n❌ *ĐÃ TỪ CHỐI*", parse_mode="Markdown")
        except: pass
    store.pop(key, None)

# ====== TẠO FILE ======
def tao_phieu_vattu(d):
    wb = load_workbook("mau_vat_tu.xlsx")
    ws = wb.active
    ngay = datetime.now()
    for row in ws.iter_rows():
        for cell in row:
            if cell.value and "Phòng/Bộ phận" in str(cell.value):
                cell.value = f"Đề nghị cung cấp cho Phòng/Bộ phận: {d['bo_phan']} một số vật tư:"
    ws['H31'] = f"Đà Nẵng, ngày {ngay.day} tháng {ngay.month} năm {ngay.year}"
    for i, vt in enumerate(d["vat_tu"][:10]):
        row = 16 + i
        ws.cell(row=row, column=2).value = vt.get("ma_vat_tu", "")
        ws.cell(row=row, column=3).value = vt.get("ten_vat_tu", "")
        ws.cell(row=row, column=4).value = vt.get("dvt", "")
        ws.cell(row=row, column=5).value = vt.get("nhu_cau", "")
        ws.cell(row=row, column=6).value = vt.get("ton_kho", "0")
        ws.cell(row=row, column=7).value = vt.get("so_luong_mua", "")
        ws.cell(row=row, column=8).value = vt.get("xuat_xu", "")
        ws.cell(row=row, column=9).value = vt.get("muc_dich", "")
        ws.cell(row=row, column=10).value = vt.get("ngay_giao", "")
    ten_file = f"YeuCauVatTu_{d['id']}.xlsx"
    wb.save(ten_file)
    return ten_file

def tao_phieu_deuxuat(d):
    doc = Document()
    doc.add_heading("CÔNG TY TNHH MTV CORNER", 0)
    doc.add_heading("PHIẾU ĐỀ XUẤT (BM/24/CORNER)", 1)
    doc.add_paragraph("Kính gửi: BAN LÃNH ĐẠO CÔNG TY TNHH MTV CORNER")
    doc.add_paragraph(f"Tôi tên là: {d['ho_ten']}    Bộ phận: {d['bo_phan']}")
    doc.add_paragraph("Nội dung đề xuất:")
    doc.add_paragraph(d['noi_dung'])
    ngay = datetime.now()
    doc.add_paragraph(f"\nĐà Nẵng, ngày {ngay.day} tháng {ngay.month} năm {ngay.year}")
    ten_file = f"PhieuDeXuat_{d['id']}.docx"
    doc.save(ten_file)
    return ten_file

def tao_phieu_thanhtoan(d):
    doc = Document()
    doc.add_heading("CÔNG TY TNHH MTV CORNER", 0)
    doc.add_heading("ĐỀ NGHỊ THANH TOÁN (BM/06/CORNER)", 1)
    doc.add_paragraph(f"Họ và tên: {d['ho_ten']}    Chức danh: {d['chuc_danh']}")
    doc.add_paragraph(f"Bộ phận: {d['bo_phan']}")
    doc.add_paragraph(f"Nội dung: {d['noi_dung']}")
    table = doc.add_table(rows=3, cols=5)
    table.style = "Table Grid"
    for i, h in enumerate(["STT", "Số HĐ", "Ngày HĐ", "Nội dung", "Số tiền"]):
        table.cell(0, i).text = h
    table.cell(1, 0).text = "1"
    table.cell(1, 1).text = d['so_hd']
    table.cell(1, 2).text = d['ngay_hd']
    table.cell(1, 3).text = d['noi_dung']
    table.cell(1, 4).text = d['so_tien']
    table.cell(2, 3).text = "Tổng"
    table.cell(2, 4).text = d['so_tien']
    doc.add_paragraph(f"\nHình thức: {d['hinh_thuc']}")
    if d['so_tk']:
        doc.add_paragraph(f"STK: {d['so_tk']} | {d['ngan_hang']} | {d['chu_tk']}")
    ngay = datetime.now()
    doc.add_paragraph(f"\nĐà Nẵng, ngày {ngay.day} tháng {ngay.month} năm {ngay.year}")
    ten_file = f"DeNghiThanhToan_{d['id']}.docx"
    doc.save(ten_file)
    return ten_file

# ====== HỦY ======
async def huy(update: Update, context: ContextTypes.DEFAULT_TYPE):
    uid = update.effective_user.id
    for key in [uid, f"dx_{uid}", f"tt_{uid}", f"excel_{uid}", f"hinh_{uid}"]:
        store.pop(key, None)
    await update.message.reply_text("❌ Đã hủy. Gõ /start để bắt đầu lại.")
    return ConversationHandler.END

def main():
    app = Application.builder().token(BOT_TOKEN).build()

    vt_conv = ConversationHandler(
        entry_points=[CommandHandler("vattu", vattu_start), CallbackQueryHandler(vattu_start, pattern="^menu_vattu$")],
        states={
            VT_BO_PHAN: [MessageHandler(filters.TEXT & ~filters.COMMAND, vt_bo_phan)],
            VT_MA: [MessageHandler(filters.TEXT & ~filters.COMMAND, vt_ma)],
            VT_TEN: [MessageHandler(filters.TEXT & ~filters.COMMAND, vt_ten)],
            VT_DVT: [MessageHandler(filters.TEXT & ~filters.COMMAND, vt_dvt)],
            VT_SO_LUONG: [MessageHandler(filters.TEXT & ~filters.COMMAND, vt_so_luong)],
            VT_MUC_DICH: [MessageHandler(filters.TEXT & ~filters.COMMAND, vt_muc_dich)],
            VT_THEM: [CallbackQueryHandler(vt_them_callback, pattern="^vt_(them|gui)$")],
            VT_XAC_NHAN: [CallbackQueryHandler(vt_xacnhan_callback, pattern="^vt_(xacnhan|huy)$")],
        },
        fallbacks=[CommandHandler("huy", huy)],
    )

    excel_conv = ConversationHandler(
        entry_points=[CommandHandler("excel", excel_start), CallbackQueryHandler(excel_start, pattern="^menu_excel$")],
        states={
            EXCEL_BO_PHAN: [MessageHandler(filters.TEXT & ~filters.COMMAND, excel_bo_phan)],
            EXCEL_XAC_NHAN: [
                MessageHandler(filters.Document.ALL, excel_nhan_file),
                CallbackQueryHandler(excel_xacnhan_callback, pattern="^excel_(ok|huy)"),
            ],
        },
        fallbacks=[CommandHandler("huy", huy)],
    )

    hinh_conv = ConversationHandler(
        entry_points=[CommandHandler("hinh", hinh_start), CallbackQueryHandler(hinh_start, pattern="^menu_hinh$")],
        states={
            HINH_BO_PHAN: [MessageHandler(filters.TEXT & ~filters.COMMAND, hinh_bo_phan)],
            HINH_XAC_NHAN: [
                MessageHandler(filters.PHOTO, hinh_nhan_anh),
                CallbackQueryHandler(hinh_xacnhan_callback, pattern="^hinh_(ok|huy)"),
            ],
        },
        fallbacks=[CommandHandler("huy", huy)],
    )

    dx_conv = ConversationHandler(
        entry_points=[CommandHandler("deuxuat", deuxuat_start), CallbackQueryHandler(deuxuat_start, pattern="^menu_deuxuat$")],
        states={
            DX_BO_PHAN: [MessageHandler(filters.TEXT & ~filters.COMMAND, dx_bo_phan)],
            DX_NOI_DUNG: [MessageHandler(filters.TEXT & ~filters.COMMAND, dx_noi_dung)],
            DX_XAC_NHAN: [CallbackQueryHandler(dx_xacnhan_callback, pattern="^dx_(xacnhan|huy)")],
        },
        fallbacks=[CommandHandler("huy", huy)],
    )

    tt_conv = ConversationHandler(
        entry_points=[CommandHandler("thanhtoan", thanhtoan_start), CallbackQueryHandler(thanhtoan_start, pattern="^menu_thanhtoan$")],
        states={
            TT_CHUC_DANH: [MessageHandler(filters.TEXT & ~filters.COMMAND, tt_chuc_danh)],
            TT_BO_PHAN: [MessageHandler(filters.TEXT & ~filters.COMMAND, tt_bo_phan)],
            TT_NOI_DUNG: [MessageHandler(filters.TEXT & ~filters.COMMAND, tt_noi_dung)],
            TT_SO_HD: [MessageHandler(filters.TEXT & ~filters.COMMAND, tt_so_hd)],
            TT_NGAY_HD: [MessageHandler(filters.TEXT & ~filters.COMMAND, tt_ngay_hd)],
            TT_SO_TIEN: [MessageHandler(filters.TEXT & ~filters.COMMAND, tt_so_tien)],
            TT_HINH_THUC: [CallbackQueryHandler(tt_hinh_thuc_callback, pattern="^tt_ht_")],
            TT_SO_TK: [MessageHandler(filters.TEXT & ~filters.COMMAND, tt_so_tk)],
            TT_NGAN_HANG: [MessageHandler(filters.TEXT & ~filters.COMMAND, tt_ngan_hang)],
            TT_CHU_TK: [MessageHandler(filters.TEXT & ~filters.COMMAND, tt_chu_tk)],
            TT_XAC_NHAN: [CallbackQueryHandler(tt_xacnhan_callback, pattern="^tt_(xacnhan|huy)")],
        },
        fallbacks=[CommandHandler("huy", huy)],
    )

    app.add_handler(CommandHandler("start", start))
    app.add_handler(vt_conv)
    app.add_handler(excel_conv)
    app.add_handler(hinh_conv)
    app.add_handler(dx_conv)
    app.add_handler(tt_conv)
    app.add_handler(CallbackQueryHandler(xu_ly_duyet, pattern="^(duyet_|tuchoi_)"))

    t = threading.Thread(target=run_health_server, daemon=True)
    t.start()
    logger.info("Bot đang chạy...")
    app.run_polling(allowed_updates=Update.ALL_TYPES)

if __name__ == "__main__":
    main()
